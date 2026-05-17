from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = Path(
    r"C:\Users\ThinkPad\Downloads\Laporan Proyek Website MI Miftahul Ulum Jipo - Revisi Akademik.docx"
)


AUTHORS = [
    "Moh. Azril Muzaki (241101004)",
    "M. Fikri Asyam Jauhary (241101008)",
    "Ahmad Khoiril Mansur Supriadi (241101011)",
    "Firzan Hanafia Asroni (241101015)",
    "Andik Prasetyo (241101016)",
    "Soleh (241101022)",
    "Lailatul Khoiriyah (241101025)",
    "Muna Rohmatun Nazila (241101113)",
]


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[index],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[index], value, align=align)
    set_table_borders(table)
    doc.add_paragraph()


def set_default_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_cover_page(doc):
    for _ in range(4):
        doc.add_paragraph()

    title_lines = [
        "LAPORAN PROYEK WEBSITE SEKOLAH",
        "MI MIFTAHUL ULUM JIPO",
    ]
    for line in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Disusun Oleh:")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)

    for author in AUTHORS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)

    for _ in range(7):
        doc.add_paragraph()

    footer_lines = [
        "PROGRAM STUDI TEKNIK INFORMATIKA",
        "FAKULTAS SAINS DAN TEKNOLOGI",
        "UNIVERSITAS NAHDLATUL ULAMA SUNAN GIRI",
        "2026",
    ]
    for line in footer_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)

    doc.add_page_break()


def add_chapter_title(doc, bab, title):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    run1 = p1.add_run(bab)
    run1.bold = True
    run1.font.name = "Times New Roman"
    run1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run1.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run(title)
    run2.bold = True
    run2.font.name = "Times New Roman"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run2.font.size = Pt(14)


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_minor_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_note_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_page_break_section(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_document():
    doc = Document()
    set_default_styles(doc)
    add_cover_page(doc)

    add_chapter_title(doc, "BAB I", "PENDAHULUAN")
    add_subheading(doc, "1.1 Latar Belakang")
    add_body_paragraph(
        doc,
        "Perkembangan teknologi informasi telah mendorong lembaga pendidikan untuk "
        "melakukan transformasi digital dalam berbagai aspek layanan. Website sekolah "
        "tidak lagi dipandang sekadar sebagai media pelengkap, melainkan sebagai sarana "
        "resmi untuk menyampaikan informasi, memperkuat citra institusi, dan membangun "
        "komunikasi yang lebih efektif dengan peserta didik, wali murid, serta masyarakat luas. "
        "Dalam konteks persaingan lembaga pendidikan yang semakin dinamis, keberadaan website "
        "resmi juga menjadi indikator profesionalitas dan kesiapan sekolah dalam menjawab kebutuhan zaman."
    )
    add_body_paragraph(
        doc,
        "Madrasah Ibtidaiyah Miftahul Ulum Jipo merupakan salah satu lembaga pendidikan dasar "
        "yang memiliki peran penting dalam pembinaan generasi muda di lingkungan Desa Jipo, "
        "Kecamatan Kepohbaru, Kabupaten Bojonegoro. Sebagai lembaga pendidikan berbasis nilai-nilai "
        "Islam, madrasah ini membutuhkan media publikasi yang mampu menampilkan profil sekolah, "
        "program unggulan, informasi akademik, prestasi, serta kegiatan kelembagaan secara terstruktur "
        "dan mudah diakses. Sebelum proyek ini dilaksanakan, penyebaran informasi masih mengandalkan "
        "cara konvensional dan media sosial, sehingga informasi sering tersebar tidak terpusat, "
        "sulit diarsipkan, dan kurang optimal bagi masyarakat yang membutuhkan referensi resmi."
    )
    add_body_paragraph(
        doc,
        "Permasalahan tersebut menimbulkan kebutuhan akan suatu sistem informasi berbasis web "
        "yang dapat berfungsi sebagai pusat informasi digital sekolah. Melalui website, sekolah "
        "dapat menyajikan identitas kelembagaan secara lebih kredibel, mempermudah calon wali murid "
        "dalam memperoleh informasi PPDB, serta memberi ruang bagi admin untuk memperbarui konten "
        "tanpa harus melakukan perubahan kode program secara langsung. Oleh karena itu, pengembangan "
        "website MI Miftahul Ulum Jipo menjadi relevan sebagai bentuk implementasi digitalisasi sekolah "
        "yang terukur, aplikatif, dan bermanfaat bagi pengelolaan informasi pendidikan."
    )
    add_body_paragraph(
        doc,
        "Selain menjawab kebutuhan praktis sekolah, proyek ini juga memiliki nilai akademik karena "
        "mengintegrasikan konsep rekayasa perangkat lunak, perancangan basis data, pengembangan web "
        "dengan framework Laravel, serta penerapan pengujian fungsional. Dengan demikian, proyek ini "
        "tidak hanya menghasilkan sebuah website, tetapi juga mendokumentasikan proses pengembangan sistem "
        "informasi secara ilmiah dan sistematis."
    )

    add_subheading(doc, "1.2 Tujuan Proyek")
    add_body_paragraph(
        doc,
        "Secara umum, proyek ini bertujuan untuk merancang dan membangun website resmi MI Miftahul "
        "Ulum Jipo sebagai media informasi digital yang mudah diakses, informatif, dan mudah dikelola. "
        "Secara lebih rinci, tujuan yang hendak dicapai adalah sebagai berikut."
    )
    add_numbered_list(
        doc,
        [
            "Membangun website sekolah berbasis Laravel yang mampu menampilkan informasi profil, akademik, berita, galeri, prestasi, kontak, dan PPDB secara terintegrasi.",
            "Menyediakan panel admin agar pihak sekolah dapat melakukan pengelolaan data secara mandiri melalui fitur tambah, lihat, ubah, dan hapus (CRUD).",
            "Mempermudah masyarakat dan calon wali murid dalam memperoleh informasi resmi sekolah tanpa harus datang langsung ke lokasi.",
            "Mendukung proses digitalisasi layanan sekolah, khususnya pada pengelolaan konten dan pendaftaran peserta didik baru secara daring.",
            "Menerapkan konsep pengembangan sistem berbasis framework modern agar aplikasi lebih terstruktur, aman, dan mudah dikembangkan di masa mendatang.",
        ],
    )

    add_subheading(doc, "1.3 Manfaat Proyek")
    add_body_paragraph(
        doc,
        "Manfaat proyek ini dapat ditinjau dari beberapa pihak yang terlibat. Bagi sekolah, website "
        "berfungsi sebagai media resmi yang meningkatkan kredibilitas institusi, mempermudah publikasi "
        "informasi, dan membantu dokumentasi kegiatan secara terpusat. Keberadaan sistem admin juga "
        "membuat pengelolaan konten menjadi lebih efisien karena pembaruan informasi dapat dilakukan "
        "melalui antarmuka yang mudah dipahami."
    )
    add_body_paragraph(
        doc,
        "Bagi masyarakat dan wali murid, website memberikan kemudahan dalam mengakses informasi kapan "
        "saja dan dari mana saja. Informasi seperti profil sekolah, tenaga pendidik, berita kegiatan, "
        "prestasi, galeri, serta informasi PPDB dapat diperoleh secara cepat melalui perangkat komputer "
        "maupun telepon pintar. Hal ini sangat membantu masyarakat yang membutuhkan sumber informasi yang "
        "lebih resmi dan terdokumentasi dengan baik."
    )
    add_body_paragraph(
        doc,
        "Bagi pengembang dan penulis laporan, proyek ini menjadi sarana penerapan kompetensi di bidang "
        "pemrograman web, pemodelan sistem, manajemen basis data, serta dokumentasi proyek perangkat lunak. "
        "Secara lebih luas, proyek ini juga menunjukkan bahwa digitalisasi sekolah dasar berbasis web dapat "
        "memberikan manfaat nyata bagi efisiensi pelayanan informasi dan peningkatan citra lembaga pendidikan."
    )

    add_page_break_section(doc)

    add_chapter_title(doc, "BAB II", "ANALISIS KEBUTUHAN DAN PERANCANGAN SISTEM")
    add_subheading(doc, "2.1 Kebutuhan Perangkat Lunak")
    add_body_paragraph(
        doc,
        "Pengembangan website dilakukan dengan memanfaatkan perangkat lunak yang mendukung pemrograman "
        "berbasis PHP, pengelolaan database relasional, pengolahan aset frontend, serta pengujian lokal "
        "melalui server development. Kebutuhan perangkat lunak yang digunakan pada proyek ini dirangkum "
        "pada Tabel 2.1."
    )
    add_table(
        doc,
        ["No", "Perangkat Lunak", "Spesifikasi / Versi", "Fungsi"],
        [
            ["1", "Sistem Operasi", "Windows 10/11", "Lingkungan kerja utama untuk proses pengembangan dan pengujian."],
            ["2", "Bahasa Pemrograman", "PHP 8.1 atau lebih baru", "Bahasa utama untuk membangun logika aplikasi."],
            ["3", "Framework Backend", "Laravel 10", "Mengatur arsitektur MVC, routing, middleware, validasi, dan ORM."],
            ["4", "Database Management System", "MySQL", "Menyimpan data pengguna, konten, PPDB, dan pengaturan website."],
            ["5", "Web Server Lokal", "Apache pada XAMPP", "Menjalankan aplikasi secara lokal selama tahap implementasi dan uji coba."],
            ["6", "Frontend Dasar", "HTML, CSS, JavaScript", "Menyusun struktur halaman, tampilan, dan interaksi antarmuka."],
            ["7", "Pengelolaan Aset Frontend", "Vite", "Memproses dan memuat aset CSS serta JavaScript secara modern."],
            ["8", "Framework / Utility CSS", "Tailwind CSS", "Mendukung penyusunan antarmuka responsif dan konsisten."],
            ["9", "Editor Kode", "Visual Studio Code", "Menulis, mengelola, dan menelusuri source code proyek."],
            ["10", "Package Manager", "Composer dan NPM", "Mengelola dependensi backend dan frontend."],
            ["11", "Library Pendukung", "DomPDF, PhpSpreadsheet, Spatie Sluggable, AOS, Swiper, GLightbox", "Mendukung ekspor PDF/Excel, pembentukan slug, animasi, slider, dan galeri."],
            ["12", "Browser Pengujian", "Google Chrome / Microsoft Edge", "Menjalankan dan menguji sistem dari sisi pengguna."],
        ],
    )
    add_note_paragraph(
        doc,
        "Catatan: pada rancangan konseptual, frontend dapat dijelaskan menggunakan istilah HTML, CSS, "
        "JavaScript, dan framework CSS. Akan tetapi, implementasi aktual proyek ini menggunakan Tailwind CSS "
        "yang diproses melalui Vite, sehingga penjelasan teknis laporan disesuaikan dengan kode yang benar-benar diuji."
    )

    add_subheading(doc, "2.2 Kebutuhan Perangkat Keras")
    add_body_paragraph(
        doc,
        "Selain perangkat lunak, proyek ini juga memerlukan perangkat keras yang memadai agar proses "
        "pengembangan, pengujian, dan akses sistem dapat berjalan dengan baik. Kebutuhan perangkat keras "
        "dibedakan menjadi perangkat untuk pengembang, pengguna, dan server lokal."
    )
    add_table(
        doc,
        ["No", "Komponen", "Spesifikasi Minimal", "Keterangan"],
        [
            ["1", "Laptop / PC Pengembang", "Prosesor Intel Core i3/Ryzen 3, RAM 8 GB, SSD/HDD 20 GB kosong", "Digunakan untuk coding, menjalankan XAMPP, Composer, NPM, dan browser pengujian."],
            ["2", "Perangkat Pengguna", "Smartphone Android/iOS atau PC dengan browser modern", "Digunakan untuk mengakses website publik dan melakukan pendaftaran PPDB."],
            ["3", "Koneksi Internet", "Stabil", "Diperlukan untuk instalasi dependensi, pengujian eksternal, dan publikasi website."],
            ["4", "Media Penyimpanan", "Minimal 10 GB tersedia", "Menyimpan source code, database, gambar unggahan, dan dokumen proyek."],
            ["5", "Server Lokal / Hosting", "CPU 2 Core, RAM 4 GB, storage 20 GB", "Spesifikasi awal yang cukup untuk website profil sekolah dan panel admin skala kecil-menengah."],
        ],
    )

    add_subheading(doc, "2.3 Arsitektur Sistem")
    add_body_paragraph(
        doc,
        "Website MI Miftahul Ulum Jipo dibangun menggunakan arsitektur Model-View-Controller (MVC) yang "
        "merupakan pola dasar pada framework Laravel. Pendekatan ini dipilih karena mampu memisahkan logika "
        "bisnis, tampilan, dan pengelolaan data secara jelas. Dengan pemisahan tersebut, proses pemeliharaan "
        "sistem menjadi lebih mudah, pengembangan fitur baru lebih terarah, dan potensi kesalahan akibat "
        "penumpukan kode pada satu bagian aplikasi dapat diminimalkan."
    )
    add_table(
        doc,
        ["Lapisan", "Komponen Utama", "Peran dalam Sistem"],
        [
            ["Presentasi", "Blade template, HTML, CSS, JavaScript, Tailwind CSS", "Menampilkan halaman publik dan halaman admin yang dapat diakses pengguna."],
            ["Aplikasi", "Route, middleware, controller, validasi, autentikasi", "Memproses permintaan pengguna, menerapkan aturan sistem, dan menghubungkan tampilan dengan data."],
            ["Data", "Model Eloquent, migration, MySQL", "Menyimpan, mengambil, dan memelihara konsistensi data aplikasi."],
            ["Infrastruktur", "Apache, PHP runtime, storage publik", "Menjalankan aplikasi dan menyajikan file unggahan seperti gambar berita, galeri, dan foto PPDB."],
        ],
    )
    add_body_paragraph(
        doc,
        "Secara konseptual, alur arsitektur sistem dapat dijelaskan sebagai berikut: pengguna mengakses "
        "website melalui browser, kemudian permintaan diterima oleh web server Apache pada lingkungan "
        "XAMPP. Selanjutnya Laravel memproses request melalui route dan middleware yang relevan. Controller "
        "akan menjalankan logika aplikasi, mengambil atau menyimpan data melalui model Eloquent ke database "
        "MySQL, lalu mengembalikan hasil dalam bentuk halaman Blade kepada pengguna. Alur ini menjadikan "
        "sistem lebih terstruktur dan konsisten dalam menangani interaksi frontend maupun backend."
    )

    add_subheading(doc, "2.4 Struktur Navigasi / Sitemap")
    add_body_paragraph(
        doc,
        "Struktur navigasi website dirancang agar sederhana, informatif, dan mudah dipahami oleh berbagai "
        "jenis pengguna. Pada sisi publik, menu difokuskan pada kemudahan akses informasi sekolah. Sementara "
        "itu, pada sisi admin, navigasi disusun berdasarkan modul pengelolaan konten dan layanan operasional."
    )
    add_table(
        doc,
        ["No", "Bagian", "Menu / Modul", "Fungsi Utama"],
        [
            ["1", "Publik", "Beranda", "Menampilkan slider, sambutan kepala madrasah, statistik, visi misi, program unggulan, berita, pengumuman, dan galeri terbaru."],
            ["2", "Publik", "Profil Madrasah", "Menjelaskan sejarah, visi, misi, akreditasi, dan struktur organisasi sekolah."],
            ["3", "Publik", "Guru & Staff", "Menampilkan data tenaga pendidik dan kependidikan."],
            ["4", "Publik", "Akademik", "Menyajikan informasi kegiatan akademik dan ekstrakurikuler."],
            ["5", "Publik", "Berita", "Menampilkan daftar berita sekolah, kategori berita, dan detail berita."],
            ["6", "Publik", "Galeri", "Menampilkan album kegiatan dan dokumentasi foto."],
            ["7", "Publik", "Prestasi", "Menyajikan data prestasi siswa dan sekolah."],
            ["8", "Publik", "PPDB", "Menampilkan syarat, jadwal, formulir pendaftaran, dan cek status pendaftaran."],
            ["9", "Publik", "Kontak", "Menyediakan informasi alamat, email, WhatsApp, peta lokasi, dan formulir pesan."],
            ["10", "Admin", "Dashboard", "Menampilkan ringkasan data utama seperti jumlah guru, berita, PPDB, dan pesan baru."],
            ["11", "Admin", "Slider, Program, Berita, Pengumuman", "Mengelola konten utama yang tampil pada halaman publik."],
            ["12", "Admin", "Album, Galeri, Guru, Struktur Organisasi, Prestasi, Ekskul", "Mengelola data dokumentasi, SDM, dan kegiatan sekolah."],
            ["13", "Admin", "PPDB dan Kontak", "Memantau pendaftar baru, memperbarui status, meninjau pesan, serta mengekspor data."],
            ["14", "Admin", "Pengaturan dan Users", "Mengubah profil sekolah global dan mengelola akun pengguna admin."],
        ],
    )

    add_subheading(doc, "2.5 Use Case Sistem")
    add_body_paragraph(
        doc,
        "Use case sistem digunakan untuk menggambarkan hubungan antara aktor dengan fungsi-fungsi utama "
        "pada aplikasi. Pada website ini terdapat tiga aktor utama, yaitu pengunjung umum, calon wali murid, "
        "dan admin sekolah. Interaksi tiap aktor dirangkum pada Tabel 2.4."
    )
    add_table(
        doc,
        ["No", "Aktor", "Use Case", "Deskripsi"],
        [
            ["1", "Pengunjung", "Melihat beranda", "Pengunjung membuka halaman utama dan memperoleh ringkasan informasi sekolah."],
            ["2", "Pengunjung", "Membaca profil sekolah", "Pengunjung melihat sejarah, visi, misi, dan struktur organisasi."],
            ["3", "Pengunjung", "Mengakses akademik, guru, prestasi, galeri, dan berita", "Pengunjung menelusuri konten sekolah sesuai kebutuhan informasi."],
            ["4", "Pengunjung", "Mengirim pesan melalui kontak", "Pengunjung mengisi formulir pesan yang kemudian tersimpan pada panel admin."],
            ["5", "Calon wali murid", "Mendaftar PPDB online", "Calon pendaftar mengisi formulir biodata dan mengirim data pendaftaran."],
            ["6", "Calon wali murid", "Mengecek status PPDB", "Calon pendaftar memasukkan nomor pendaftaran untuk melihat status verifikasi."],
            ["7", "Admin", "Login ke sistem", "Admin memasukkan email dan password untuk mengakses dashboard."],
            ["8", "Admin", "Mengelola konten website", "Admin melakukan CRUD pada berita, pengumuman, galeri, guru, program unggulan, dan modul lainnya."],
            ["9", "Admin", "Memproses PPDB", "Admin meninjau data pendaftar, memperbarui status, dan menambahkan catatan bila diperlukan."],
            ["10", "Admin", "Mengelola pengaturan global", "Admin memperbarui identitas sekolah, statistik, visi misi, kontak, dan aset seperti logo serta favicon."],
        ],
    )

    add_subheading(doc, "2.6 Flowchart Sistem")
    add_body_paragraph(
        doc,
        "Flowchart sistem menggambarkan alur kerja utama aplikasi dari sisi pengguna maupun admin. Pada laporan ini, "
        "flowchart dijelaskan dalam bentuk urutan proses agar mudah dipahami dan dapat dikembangkan kembali menjadi "
        "diagram visual apabila diperlukan pada presentasi seminar."
    )
    add_minor_heading(doc, "a. Alur Pengunjung Mengakses Informasi")
    add_numbered_list(
        doc,
        [
            "Pengunjung membuka alamat website melalui browser.",
            "Sistem menampilkan halaman beranda beserta data slider, program unggulan, berita, pengumuman, dan galeri terbaru.",
            "Pengunjung memilih menu yang dibutuhkan, misalnya Profil, Akademik, Berita, Galeri, Prestasi, atau Kontak.",
            "Laravel memproses request melalui route dan controller yang sesuai.",
            "Data diambil dari database MySQL melalui model Eloquent.",
            "Halaman Blade dirender dan hasilnya dikirim kembali ke browser.",
            "Pengunjung menerima informasi yang diminta dan dapat melanjutkan navigasi ke halaman lain.",
        ],
    )
    add_minor_heading(doc, "b. Alur Pendaftaran PPDB Online")
    add_numbered_list(
        doc,
        [
            "Calon wali murid membuka halaman PPDB.",
            "Sistem menampilkan jadwal pendaftaran, syarat, dan formulir PPDB.",
            "Pengguna mengisi data siswa dan data orang tua sesuai kolom yang tersedia.",
            "Sistem melakukan validasi terhadap data wajib, format email, jenis kelamin, tanggal lahir, dan file foto bila diunggah.",
            "Apabila data tidak valid, sistem menampilkan pesan kesalahan pada form.",
            "Apabila data valid, sistem menyimpan data ke tabel PPDB dan membuat nomor pendaftaran otomatis.",
            "Pengguna diarahkan ke halaman cek status dan memperoleh notifikasi bahwa pendaftaran berhasil dikirim.",
        ],
    )
    add_minor_heading(doc, "c. Alur Admin Mengelola Website")
    add_numbered_list(
        doc,
        [
            "Admin membuka halaman login dan memasukkan email serta password.",
            "Laravel memverifikasi kredensial menggunakan mekanisme autentikasi bawaan.",
            "Jika login berhasil, session diregenerasi dan admin diarahkan ke dashboard.",
            "Admin memilih modul yang ingin dikelola, seperti berita, galeri, pengumuman, guru, prestasi, atau PPDB.",
            "Admin melakukan operasi CRUD melalui form yang telah disediakan.",
            "Sistem memvalidasi input, menyimpan perubahan ke database, dan mengelola file unggahan pada storage publik.",
            "Konten yang telah diperbarui langsung dapat ditampilkan pada halaman publik website.",
        ],
    )

    add_subheading(doc, "2.7 Perancangan Database")
    add_body_paragraph(
        doc,
        "Perancangan database menggunakan MySQL dengan pendekatan basis data relasional. Struktur tabel disusun "
        "untuk mendukung pengelolaan konten website, layanan PPDB, komunikasi pengunjung, serta pengaturan sekolah. "
        "Dengan desain ini, data dapat dipisahkan sesuai fungsi masing-masing sehingga memudahkan proses pemeliharaan, "
        "pengembangan, dan pencarian data."
    )
    add_table(
        doc,
        ["No", "Nama Tabel", "Fungsi Utama", "Field Penting / Relasi"],
        [
            ["1", "users", "Menyimpan akun admin sistem.", "name, email, password, role"],
            ["2", "pengaturan", "Menyimpan pengaturan global website.", "key, value"],
            ["3", "sliders", "Menyimpan banner pada halaman beranda.", "image, caption, link, urutan, is_active"],
            ["4", "program_unggulan", "Menyimpan data program unggulan sekolah.", "judul, deskripsi, gambar, warna, urutan, is_active"],
            ["5", "kategori_berita", "Menyimpan kategori artikel berita.", "nama, slug"],
            ["6", "berita", "Menyimpan artikel berita sekolah.", "judul, slug, konten, gambar, kategori_id, user_id, published_at"],
            ["7", "pengumuman", "Menyimpan pengumuman resmi sekolah.", "judul, konten, tanggal_mulai, tanggal_selesai, is_active"],
            ["8", "album_galeri", "Menyimpan album dokumentasi kegiatan.", "nama, deskripsi, cover, tanggal"],
            ["9", "galeri", "Menyimpan foto per album galeri.", "album_id, gambar, caption, urutan"],
            ["10", "guru", "Menyimpan data guru dan staff.", "nama, nip, jabatan, mata_pelajaran, pendidikan, email, bio"],
            ["11", "struktur_org", "Menyimpan susunan struktur organisasi sekolah.", "nama, jabatan, foto, urutan"],
            ["12", "prestasi", "Menyimpan data prestasi siswa atau sekolah.", "nama_prestasi, kategori, tingkat, tahun, penyelenggara"],
            ["13", "ekskul", "Menyimpan data kegiatan ekstrakurikuler.", "nama, deskripsi, pembina, jadwal, is_active"],
            ["14", "ppdb", "Menyimpan data pendaftaran peserta didik baru.", "no_pendaftaran, nama_siswa, orang_tua, alamat, telepon, status, tahun_ajaran"],
            ["15", "kontak_pesan", "Menyimpan pesan dari formulir kontak.", "nama, email, telepon, subjek, pesan, is_read"],
        ],
    )
    add_body_paragraph(
        doc,
        "Relasi utama pada database ini meliputi hubungan satu-ke-banyak antara tabel kategori_berita dengan berita, "
        "users dengan berita, serta album_galeri dengan galeri. Tabel pengaturan menggunakan pendekatan key-value "
        "agar data profil sekolah, statistik, dan elemen global website dapat dikelola secara fleksibel. Sementara itu, "
        "tabel ppdb dan kontak_pesan berperan sebagai tabel transaksi yang menangkap interaksi pengguna dari sisi publik."
    )

    add_page_break_section(doc)

    add_chapter_title(doc, "BAB III", "IMPLEMENTASI DAN CARA KERJA SISTEM")
    add_subheading(doc, "3.1 Implementasi Antarmuka")
    add_body_paragraph(
        doc,
        "Implementasi antarmuka website difokuskan pada kemudahan navigasi, keterbacaan informasi, dan kesesuaian "
        "dengan identitas lembaga pendidikan. Halaman publik disusun dengan layout yang konsisten melalui Blade layout, "
        "sehingga elemen seperti navbar, footer, dan identitas visual dapat digunakan kembali pada seluruh halaman. "
        "Pendekatan ini membantu menjaga konsistensi tampilan serta mempercepat proses pengembangan."
    )
    add_table(
        doc,
        ["No", "Halaman", "Implementasi Antarmuka", "Fungsi"],
        [
            ["1", "Beranda", "Slider utama, statistik sekolah, sambutan kepala madrasah, visi misi, program unggulan, berita, pengumuman, galeri, dan tombol CTA PPDB.", "Menjadi pusat informasi awal bagi pengunjung."],
            ["2", "Profil", "Halaman sejarah, visi misi, akreditasi, dan struktur organisasi.", "Menjelaskan identitas serta karakter kelembagaan sekolah."],
            ["3", "Akademik", "Informasi kegiatan akademik dan ekstrakurikuler.", "Memberikan gambaran aktivitas pembelajaran dan pengembangan siswa."],
            ["4", "Guru & Staff", "Kartu data guru dan tenaga kependidikan.", "Menampilkan SDM sekolah secara ringkas dan informatif."],
            ["5", "Berita", "Daftar berita, kategori, dan halaman detail berdasarkan slug.", "Menyebarkan informasi kegiatan dan publikasi sekolah."],
            ["6", "Galeri", "Daftar album serta foto dokumentasi yang dapat dibuka dalam lightbox.", "Mendukung dokumentasi visual kegiatan sekolah."],
            ["7", "Prestasi", "Daftar capaian siswa atau lembaga.", "Menunjukkan kualitas dan perkembangan prestasi madrasah."],
            ["8", "PPDB", "Jadwal, syarat, formulir pendaftaran, dan cek status.", "Mendukung layanan pendaftaran peserta didik baru secara daring."],
            ["9", "Kontak", "Informasi alamat, peta lokasi, dan formulir pesan.", "Mempermudah komunikasi antara sekolah dengan masyarakat."],
        ],
    )
    add_body_paragraph(
        doc,
        "Dari sisi pengalaman pengguna, antarmuka website juga memanfaatkan elemen visual modern seperti slider, "
        "animasi ringan, dan tampilan kartu untuk menyajikan informasi secara lebih menarik. Implementasi JavaScript "
        "digunakan untuk mendukung navigasi mobile, animasi penghitung statistik, smooth scroll, slider hero, lightbox "
        "galeri, dan notifikasi antarmuka. Dengan demikian, pengguna memperoleh pengalaman akses yang lebih interaktif "
        "tanpa mengurangi kejelasan isi informasi."
    )

    add_subheading(doc, "3.2 Implementasi Backend Laravel")
    add_body_paragraph(
        doc,
        "Backend sistem dibangun menggunakan framework Laravel 10 dengan pola arsitektur MVC. Pada lapisan route, "
        "pengelompokan jalur akses dilakukan secara terstruktur, misalnya route publik untuk profil, berita, galeri, "
        "prestasi, PPDB, dan kontak; serta route khusus admin yang diletakkan dalam prefix admin dan dilindungi oleh "
        "middleware autentikasi. Pengorganisasian route ini memudahkan pemeliharaan sekaligus memperjelas pemisahan "
        "antara layanan publik dan layanan administratif."
    )
    add_body_paragraph(
        doc,
        "Pada lapisan controller, Laravel digunakan untuk menangani logika bisnis utama. HomeController mengambil data "
        "slider, program unggulan, berita, pengumuman, galeri, dan pengaturan untuk ditampilkan pada beranda. "
        "PpdbController menangani validasi form pendaftaran, unggah foto siswa, pembuatan nomor pendaftaran otomatis, "
        "dan pengecekan status pendaftaran. KontakController memvalidasi form pesan dan menyimpannya ke database, "
        "sementara controller admin mengelola berbagai modul CRUD seperti berita, galeri, guru, prestasi, program, "
        "pengumuman, kontak, dan pengaturan website."
    )
    add_body_paragraph(
        doc,
        "Pada lapisan model, Eloquent ORM mempermudah proses interaksi dengan tabel-tabel MySQL melalui representasi objek. "
        "Sebagai contoh, model Berita memanfaatkan package Spatie Sluggable untuk membentuk slug artikel secara otomatis "
        "berdasarkan judul, sedangkan model Ppdb memiliki metode pembangkit nomor pendaftaran berformat terstruktur. "
        "Pendekatan ini membuat pengelolaan query lebih rapi, aman, dan mudah dibaca dibandingkan penulisan SQL mentah "
        "secara langsung di dalam tampilan."
    )
    add_body_paragraph(
        doc,
        "Laravel juga mendukung pengelolaan file unggahan melalui storage publik. Gambar berita, slider, galeri, dan foto "
        "pendaftar disimpan secara terorganisasi pada direktori storage/app/public, kemudian diakses melalui symbolic link "
        "ke public/storage. Selain itu, proyek ini memanfaatkan package DomPDF untuk ekspor data PPDB ke PDF dan "
        "PhpSpreadsheet untuk ekspor ke Excel. Integrasi fitur-fitur ini menunjukkan bahwa Laravel tidak hanya berfungsi "
        "sebagai framework dasar, tetapi juga sebagai fondasi pengembangan sistem yang lebih komprehensif."
    )

    add_subheading(doc, "3.3 Cara Kerja Fitur Utama")
    add_body_paragraph(
        doc,
        "Cara kerja sistem dapat dipahami melalui fitur-fitur utama yang menjadi kebutuhan sekolah. Setiap fitur dirancang "
        "agar saling terhubung, mudah digunakan, dan relevan dengan proses penyampaian informasi serta pelayanan sekolah."
    )
    add_numbered_list(
        doc,
        [
            "Beranda dinamis memuat data dari beberapa tabel sekaligus, seperti slider, program unggulan, berita, pengumuman, galeri, dan pengaturan. Dengan demikian, halaman utama selalu menampilkan informasi terbaru yang diinput melalui panel admin.",
            "Fitur berita memanfaatkan slug artikel agar alamat URL lebih rapi, mudah dibaca, dan lebih baik dari sisi optimasi mesin pencari. Pengelompokan berita berdasarkan kategori juga membantu pengguna menemukan konten sesuai topik.",
            "Fitur galeri menerapkan struktur album dan foto. Admin terlebih dahulu membuat album, kemudian mengunggah foto ke album terkait. Pengunjung dapat membuka detail album untuk melihat dokumentasi kegiatan secara terorganisasi.",
            "Fitur prestasi dan data guru berfungsi sebagai media publikasi capaian serta sumber daya manusia sekolah. Informasi ini penting untuk memperlihatkan kualitas lembaga kepada masyarakat.",
            "Fitur kontak memungkinkan pengunjung mengirim pesan langsung melalui form. Data pesan tidak hilang karena tersimpan pada database dan dapat dipantau kembali oleh admin melalui dashboard.",
            "Fitur PPDB online menerima data calon peserta didik, membuat nomor pendaftaran otomatis, dan menyediakan halaman cek status. Mekanisme ini membantu sekolah dalam proses pencatatan awal pendaftar secara lebih rapi dan terdokumentasi.",
        ],
    )

    add_subheading(doc, "3.4 Sistem Admin")
    add_body_paragraph(
        doc,
        "Sistem admin menjadi komponen inti dalam pengelolaan website karena hampir seluruh konten publik bersumber dari panel ini. "
        "Setelah berhasil login, admin akan diarahkan ke dashboard yang menampilkan ringkasan jumlah guru, total berita, total "
        "pendaftar PPDB, pendaftar baru, dan pesan masuk terbaru. Dashboard juga menyediakan akses cepat menuju modul yang sering digunakan."
    )
    add_table(
        doc,
        ["No", "Modul Admin", "Operasi yang Didukung", "Keterangan"],
        [
            ["1", "Berita", "Create, Read, Update, Delete", "Mengelola artikel beserta kategori, status publish, dan gambar."],
            ["2", "Pengumuman", "Create, Read, Update, Delete", "Mengelola pengumuman aktif berdasarkan rentang tanggal."],
            ["3", "Slider", "Create, Read, Update, Delete", "Mengelola banner beranda yang ditampilkan pada halaman utama."],
            ["4", "Program Unggulan", "Create, Read, Update, Delete", "Mengelola program prioritas sekolah pada homepage."],
            ["5", "Album dan Galeri", "Create, Read, Update, Delete", "Mengelola album serta foto dokumentasi kegiatan."],
            ["6", "Guru, Struktur Organisasi, Prestasi, Ekskul", "Create, Read, Update, Delete", "Mengelola data sumber daya manusia dan kegiatan sekolah."],
            ["7", "PPDB", "Read, Update Status, Export", "Meninjau pendaftar, mengubah status, dan mengekspor data ke Excel/PDF."],
            ["8", "Kontak", "Read, Mark as Read, Delete", "Membaca pesan pengunjung dan menandai status baca."],
            ["9", "Pengaturan Web", "Read, Update", "Mengelola identitas sekolah, statistik, visi misi, sejarah, logo, favicon, dan kontak."],
            ["10", "Users", "Create, Read, Update, Delete", "Mengelola akun pengguna panel admin."],
        ],
    )
    add_minor_heading(doc, "Keamanan Login dan Panel Admin")
    add_body_paragraph(
        doc,
        "Aspek keamanan sistem admin memanfaatkan mekanisme bawaan Laravel yang cukup andal untuk skala aplikasi sekolah. "
        "Pada saat login, sistem melakukan validasi email dan password terlebih dahulu. Jika autentikasi berhasil, Laravel "
        "meregenerasi session untuk mengurangi risiko session fixation, kemudian mengarahkan pengguna ke dashboard. Ketika "
        "logout dilakukan, session di-invalidasi dan token CSRF diregenerasi agar sesi sebelumnya tidak dapat digunakan kembali."
    )
    add_body_paragraph(
        doc,
        "Seluruh route admin utama dilindungi middleware auth sehingga halaman pengelolaan data tidak dapat diakses tanpa login. "
        "Di sisi form, Laravel menerapkan proteksi CSRF, sedangkan validasi sisi server digunakan untuk memastikan data yang masuk "
        "sesuai aturan, misalnya pemeriksaan format email, data wajib, tipe file gambar, dan batas ukuran unggahan. Password akun "
        "juga disimpan dalam bentuk hash, sehingga tidak tersimpan sebagai teks biasa di database."
    )
    add_body_paragraph(
        doc,
        "Meskipun demikian, untuk pengembangan lanjutan masih disarankan penambahan lapisan otorisasi berbasis peran yang lebih ketat, "
        "misalnya pemisahan hak akses antara admin biasa dan superadmin, pencatatan log aktivitas, pembatasan percobaan login, dan "
        "integrasi CAPTCHA pada form publik. Penjelasan ini penting agar laporan bersifat objektif dan tetap mempertimbangkan ruang pengembangan sistem."
    )

    add_subheading(doc, "3.5 Responsive Design")
    add_body_paragraph(
        doc,
        "Website dirancang dengan pendekatan responsive design agar dapat diakses dengan baik melalui desktop, tablet, maupun smartphone. "
        "Implementasi responsivitas dilakukan melalui pengaturan grid, flexbox, ukuran teks yang adaptif, dan perubahan susunan navigasi "
        "antara mode desktop dan mobile. Pada layar besar, menu utama ditampilkan secara horizontal, sedangkan pada layar kecil sistem "
        "mengubahnya menjadi menu mobile yang dapat dibuka dan ditutup secara interaktif."
    )
    add_body_paragraph(
        doc,
        "Komponen lain seperti kartu berita, daftar galeri, formulir PPDB, dan formulir kontak juga disusun agar kolom-kolom input tetap nyaman "
        "dibaca pada berbagai ukuran layar. Pendekatan ini penting karena mayoritas pengguna website sekolah cenderung mengakses melalui perangkat "
        "seluler. Dengan desain responsif, kualitas pengalaman pengguna tetap terjaga dan informasi sekolah dapat tersampaikan secara optimal."
    )

    add_page_break_section(doc)

    add_chapter_title(doc, "BAB IV", "PENGUJIAN DAN HASIL")
    add_subheading(doc, "4.1 Skenario Pengujian")
    add_body_paragraph(
        doc,
        "Pengujian sistem dilakukan menggunakan pendekatan black box testing, yaitu pengujian yang berfokus pada fungsi sistem tanpa melihat "
        "detail kode program secara langsung dari sudut pandang pengguna akhir. Pengujian dilakukan pada lingkungan lokal menggunakan XAMPP, "
        "PHP 8.1, MySQL, dan browser modern. Skenario uji dipilih berdasarkan fitur yang paling penting bagi operasional website sekolah."
    )
    add_table(
        doc,
        ["No", "Modul", "Skenario Uji", "Hasil yang Diharapkan"],
        [
            ["1", "Beranda", "Membuka halaman utama website.", "Seluruh komponen utama tampil tanpa error dan data termuat dengan benar."],
            ["2", "Berita", "Membuka daftar berita dan detail berita.", "Daftar berita dan detail artikel tampil sesuai kategori dan slug."],
            ["3", "Galeri", "Membuka daftar album dan foto.", "Album tampil dan foto dapat dibuka sesuai album yang dipilih."],
            ["4", "Kontak", "Mengirim form kontak dengan data valid.", "Pesan tersimpan dan notifikasi sukses ditampilkan."],
            ["5", "Kontak", "Mengirim form kontak dengan email tidak valid.", "Sistem menolak input dan menampilkan validasi kesalahan."],
            ["6", "PPDB", "Mengirim formulir PPDB dengan data lengkap.", "Data tersimpan, nomor pendaftaran terbentuk otomatis, dan notifikasi sukses tampil."],
            ["7", "PPDB", "Mengirim formulir PPDB dengan data wajib kosong.", "Sistem menolak pengiriman dan menampilkan pesan validasi."],
            ["8", "PPDB", "Mengecek status pendaftaran dengan nomor yang benar.", "Status pendaftaran tampil sesuai data pada database."],
            ["9", "Login Admin", "Login menggunakan email dan password yang benar.", "Admin berhasil masuk ke dashboard."],
            ["10", "Login Admin", "Login menggunakan password yang salah.", "Sistem menolak login dan menampilkan pesan kesalahan."],
            ["11", "Admin Berita", "Menambah dan memperbarui berita.", "Data berita tersimpan dan perubahan tampil pada daftar berita."],
            ["12", "Admin PPDB", "Mengubah status pendaftaran dan ekspor data.", "Status berubah sesuai pilihan dan file ekspor berhasil diunduh."],
        ],
    )

    add_subheading(doc, "4.2 Black Box Testing")
    add_body_paragraph(
        doc,
        "Hasil pengujian black box menunjukkan bahwa fungsi-fungsi utama sistem telah berjalan sesuai kebutuhan. Ringkasan hasil pengujian disajikan "
        "pada Tabel 4.2 sebagai berikut."
    )
    add_table(
        doc,
        ["No", "Fitur yang Diuji", "Output yang Diharapkan", "Hasil Aktual", "Status"],
        [
            ["1", "Akses Beranda", "Beranda tampil lengkap.", "Slider, statistik, program, berita, pengumuman, dan galeri berhasil tampil.", "Berhasil"],
            ["2", "Navigasi Menu Publik", "Seluruh menu dapat diakses.", "Menu Profil, Akademik, Guru, Berita, Galeri, Prestasi, Kontak, dan PPDB dapat dibuka.", "Berhasil"],
            ["3", "Detail Berita", "Artikel tampil berdasarkan slug.", "Sistem menampilkan detail berita sesuai data yang dipilih.", "Berhasil"],
            ["4", "Form Kontak Valid", "Pesan tersimpan.", "Pesan masuk tersimpan di tabel kontak_pesan dan muncul notifikasi sukses.", "Berhasil"],
            ["5", "Form Kontak Tidak Valid", "Input ditolak.", "Sistem menampilkan pesan validasi untuk email yang salah.", "Berhasil"],
            ["6", "PPDB Valid", "Nomor pendaftaran terbentuk otomatis.", "Sistem membuat nomor format PPDB-TAHUN-XXXX dan menyimpan data.", "Berhasil"],
            ["7", "PPDB Tidak Valid", "Data tidak disimpan.", "Kolom wajib yang kosong memunculkan pesan validasi dan proses dibatalkan.", "Berhasil"],
            ["8", "Cek Status PPDB", "Status tampil sesuai database.", "Data pendaftaran dan status verifikasi berhasil ditampilkan.", "Berhasil"],
            ["9", "Login Admin Valid", "Masuk ke dashboard.", "Admin diarahkan ke dashboard dan session aktif.", "Berhasil"],
            ["10", "Login Admin Tidak Valid", "Akses ditolak.", "Sistem menampilkan pesan email atau password salah.", "Berhasil"],
            ["11", "CRUD Berita", "Data dapat ditambah, diubah, dan dihapus.", "Operasi CRUD berjalan normal termasuk unggah gambar.", "Berhasil"],
            ["12", "Update Status PPDB", "Status pendaftar berubah.", "Perubahan status tersimpan dan dapat dilihat kembali pada detail pendaftar.", "Berhasil"],
            ["13", "Ekspor Data PPDB", "File Excel dan PDF terunduh.", "Ekspor data berhasil menghasilkan file unduhan.", "Berhasil"],
            ["14", "Tampilan Mobile", "Layout menyesuaikan layar kecil.", "Menu mobile, form, dan konten tetap rapi pada layar smartphone.", "Berhasil"],
        ],
    )

    add_subheading(doc, "4.3 Hasil Pengujian")
    add_body_paragraph(
        doc,
        "Berdasarkan pengujian yang telah dilakukan, sistem menunjukkan kinerja fungsional yang baik pada modul-modul utama. Seluruh menu publik "
        "dapat diakses dengan normal, tidak ditemukan tautan rusak pada jalur navigasi yang diuji, dan konten berhasil ditampilkan sesuai data yang "
        "disimpan pada database. Fitur PPDB online mampu menerima input, menghasilkan nomor pendaftaran otomatis, serta menyediakan halaman pengecekan status."
    )
    add_body_paragraph(
        doc,
        "Dari sisi admin, proses login, pengelolaan konten, pembaruan status PPDB, hingga ekspor data berjalan sebagaimana mestinya. Validasi input di sisi "
        "server juga membantu mencegah data kosong atau data berformat salah masuk ke sistem. Dengan kata lain, website ini telah memenuhi kebutuhan dasar "
        "sebagai media informasi sekolah sekaligus alat bantu administrasi digital untuk pengelolaan konten dan pendaftaran."
    )

    add_subheading(doc, "4.4 Analisis Hasil")
    add_body_paragraph(
        doc,
        "Hasil pengujian memperlihatkan bahwa pemanfaatan Laravel memberikan dampak positif terhadap keteraturan kode dan kestabilan aplikasi. Struktur MVC, "
        "validasi request, middleware autentikasi, serta ORM Eloquent membuat pengembangan fitur lebih terkontrol. Dari sisi pengguna, alur akses informasi "
        "dan layanan PPDB terasa lebih sederhana karena seluruh fitur penting telah dihimpun dalam satu platform yang terintegrasi."
    )
    add_body_paragraph(
        doc,
        "Meskipun hasil fungsional menunjukkan performa yang baik, terdapat beberapa hal yang perlu dicatat secara objektif. Pengujian pada laporan ini masih "
        "berfokus pada black box testing dan belum mencakup pengujian beban tinggi, uji keamanan penetrasi, maupun pengukuran performa server secara mendalam. "
        "Selain itu, sistem otorisasi berbasis peran masih dapat dikembangkan agar kontrol akses antar pengguna admin menjadi lebih spesifik. Dengan demikian, "
        "website sudah layak digunakan untuk kebutuhan informasi sekolah, namun tetap memiliki ruang perbaikan agar semakin kuat untuk penggunaan jangka panjang."
    )

    add_page_break_section(doc)

    add_chapter_title(doc, "BAB V", "PENUTUP")
    add_subheading(doc, "5.1 Kesimpulan")
    add_body_paragraph(
        doc,
        "Berdasarkan hasil perancangan, implementasi, dan pengujian yang telah dilakukan, dapat disimpulkan bahwa proyek pengembangan website MI Miftahul "
        "Ulum Jipo berhasil diwujudkan sebagai media informasi digital sekolah yang lebih profesional, terstruktur, dan mudah diakses. Sistem ini mampu "
        "menampilkan informasi profil sekolah, berita, galeri, prestasi, data guru, kontak, dan layanan PPDB online dalam satu platform yang terintegrasi."
    )
    add_body_paragraph(
        doc,
        "Dari sisi teknis, penggunaan Laravel terbukti mendukung pembangunan aplikasi yang lebih rapi melalui pola MVC, validasi data, autentikasi admin, "
        "pengelolaan database dengan Eloquent, serta kemudahan pengembangan modul CRUD. Kehadiran panel admin juga memberi manfaat praktis karena pihak sekolah "
        "dapat memperbarui konten secara mandiri tanpa harus selalu bergantung pada perubahan kode program. Dengan demikian, website ini tidak hanya relevan "
        "sebagai hasil proyek akademik, tetapi juga berpotensi digunakan secara nyata untuk mendukung digitalisasi layanan sekolah."
    )

    add_subheading(doc, "5.2 Saran")
    add_body_paragraph(
        doc,
        "Agar sistem dapat berkembang lebih optimal di masa mendatang, beberapa saran pengembangan yang dapat dipertimbangkan adalah sebagai berikut."
    )
    add_numbered_list(
        doc,
        [
            "Menambahkan otorisasi berbasis peran yang lebih detail, misalnya pemisahan hak akses antara superadmin, admin konten, dan admin PPDB.",
            "Mengintegrasikan notifikasi otomatis melalui email atau WhatsApp untuk pendaftar PPDB maupun pesan masuk dari pengunjung.",
            "Menambahkan fitur keamanan lanjutan seperti CAPTCHA, pembatasan percobaan login, backup data berkala, dan pencatatan log aktivitas admin.",
            "Mengembangkan modul akademik yang lebih lengkap, seperti agenda sekolah, arsip dokumen, atau sistem informasi akademik sederhana.",
            "Melakukan deployment ke hosting produksi serta pengujian performa dan keamanan yang lebih mendalam sebelum sistem digunakan secara luas.",
        ],
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(f"Dokumen berhasil dibuat: {OUTPUT_PATH}")
